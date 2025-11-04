# -*- coding: utf-8 -*-
import pandas as pd
import numpy as np
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.cluster import KMeans, AgglomerativeClustering
from sklearn.metrics import adjusted_rand_score, silhouette_score, normalized_mutual_info_score, completeness_score, homogeneity_score
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
import sys
import io

# Windows terminalinde UTF-8 encoding sorununu çöz
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# Uyarıları bastır
warnings.filterwarnings('ignore')

# DOCX rapor için Document oluştur
doc = Document()

# === DOCX YARDIMCI FONKSİYONLAR ===

def add_heading_docx(doc, text, level=1):
    """DOCX'e başlık ekle"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph_docx(doc, text, bold=False, italic=False):
    """DOCX'e paragraf ekle"""
    para = doc.add_paragraph(text)
    if bold:
        for run in para.runs:
            run.bold = True
    if italic:
        for run in para.runs:
            run.italic = True
    return para

def add_table_docx(doc, headers, rows):
    """DOCX'e tablo ekle"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Başlıkları ekle
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Satırları ekle
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    return table

def print_separator(title=""):
    """Bölüm ayırıcı yazdır"""
    print("\n" + "="*80)
    if title:
        print(f" {title}")
        print("="*80)
    else:
        print("="*80)

# === RAPOR BAŞLANGIÇ ===
print_separator("KÜMELEME (CLUSTERING) ANALİZ RAPORU")
print("\nAcoustic Features Dataset - Kümeleme Yöntemleri Karşılaştırması\n")

# DOCX başlık
add_heading_docx(doc, "KÜMELEME (CLUSTERING) ANALİZ RAPORU", level=0)
add_paragraph_docx(doc, "Acoustic Features Dataset - Kümeleme Yöntemleri Karşılaştırması")
add_paragraph_docx(doc, "")

# === 1. PROBLEM AÇIKLAMASI ===
print_separator("1. PROBLEM AÇIKLAMASI")

problem_text = """Bu çalışmada, müzik parçalarının akustik özelliklerine göre kümeleme (clustering) analizi yapılmaktadır. Kümeleme, gözetimsiz öğrenme (unsupervised learning) yöntemidir ve verideki doğal grupları keşfetmeye çalışır. Gerçek etiketler bilindiği için (angry, happy, relax, sad), kümeleme yöntemlerinin performansını gerçek etiketlerle karşılaştırarak değerlendirebiliriz. İki farklı kümeleme algoritması (K-Means ve Agglomerative Clustering) kullanılarak veri kümelere ayrılmış ve başarı metrikleri ile karşılaştırılmıştır."""

print("\n" + problem_text + "\n")

# DOCX problem açıklaması
add_heading_docx(doc, "1. PROBLEM AÇIKLAMASI", level=1)
add_paragraph_docx(doc, problem_text)

# === 2. KÜMELEME YÖNTEMLERİNİN AÇIKLAMASI ===
print_separator("2. KULLANILAN KÜMELEME YÖNTEMLERİ")

method_text_kmeans = """K-Means, en popüler kümeleme algoritmalarından biridir. Algoritma, veriyi k adet kümeye ayırmaya çalışır. İşleyişi şu şekildedir: (1) k adet rastgele merkez (centroid) seçilir, (2) Her veri noktası en yakın merkeze atanır, (3) Merkezler, atanan noktaların ortalaması olarak güncellenir, (4) Adımlar 2-3 iterasyon olarak tekrarlanır ve küme merkezleri sabitlenince durur. Avantajları: Hızlı ve ölçeklenebilir, basit ve anlaşılır. Dezavantajları: Küme sayısının önceden belirlenmesi gerekir, küresel olmayan kümelerde zorlanır."""

method_text_agg = """Agglomerative Clustering (Hiyerarşik Kümeleme), hiyerarşik bir kümeleme yöntemidir. Algoritma, her veri noktasını başlangıçta ayrı bir küme olarak ele alır ve en yakın kümeleri iteratif olarak birleştirir (bottom-up yaklaşım). Ward linkage kriteri kullanıldığında, birleştirme işlemi kümeler içi varyansı minimize edecek şekilde yapılır. Avantajları: Küme sayısını sonradan belirleme esnekliği, dendrogram ile görselleştirme imkanı. Dezavantajları: Hesaplama maliyeti yüksek (O(n³)), büyük veri setlerinde yavaş."""

print("\n2.1. K-Means Kümeleme")
print("-"*80)
print(method_text_kmeans)
print("\n2.2. Agglomerative Clustering (Hiyerarşik Kümeleme)")
print("-"*80)
print(method_text_agg + "\n")

# DOCX yöntem açıklamaları
add_heading_docx(doc, "2. KULLANILAN KÜMELEME YÖNTEMLERİ", level=1)
add_heading_docx(doc, "2.1. K-Means Kümeleme", level=2)
add_paragraph_docx(doc, method_text_kmeans)
add_heading_docx(doc, "2.2. Agglomerative Clustering (Hiyerarşik Kümeleme)", level=2)
add_paragraph_docx(doc, method_text_agg)

# === 3. VERİ YÜKLEME VE HAZIRLAMA ===
print_separator("3. VERİ YÜKLEME VE HAZIRLAMA")

try:
    # Veri setini yükle (Acoustic Features.csv kullanılıyor)
    df = pd.read_csv('Acoustic Features.csv')
    
    # Class sütununu etiketler (y_true) olarak ayır
    y_labels = df['Class'].values
    
    # Class sütunu dışındaki tüm sütunları özellikler (X) olarak al
    X = df.drop('Class', axis=1).values
    
    # Etiketleri encode et
    le = LabelEncoder()
    y_true = le.fit_transform(y_labels)
    class_names = le.classes_
    
    # Bilinen sınıf sayısı
    n_clusters = len(class_names)
    
    print(f"\nVeri Seti Yüklendi.")
    print(f"Örnek Sayısı: {X.shape[0]}, Öznitelik Sayısı: {X.shape[1]}")
    print(f"Küme Sayısı: {n_clusters}")
    print(f"Sınıflar: {class_names}\n")
    
    # Veriyi normalize et (kümeleme için kritik)
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)
    
    print("Veri normalize edildi (StandardScaler).\n")
    
    # DOCX veri bilgisi
    add_heading_docx(doc, "3. VERİ YÜKLEME VE HAZIRLAMA", level=1)
    add_paragraph_docx(doc, f"Örnek Sayısı: {X.shape[0]}, Öznitelik Sayısı: {X.shape[1]}")
    add_paragraph_docx(doc, f"Küme Sayısı: {n_clusters}")
    add_paragraph_docx(doc, f"Sınıflar: {', '.join(class_names)}")
    add_paragraph_docx(doc, "Veri StandardScaler ile normalize edilmiştir.")

    # === 4. KÜMELEME YÖNTEMLERİNİN UYGULANMASI ===
    print_separator("4. KÜMELEME YÖNTEMLERİNİN UYGULANMASI")
    
    print("--- Kümeleme Yöntemleri Uygulanıyor ---\n")
    
    # Yöntem 1: K-Means
    print("K-Means kümeleme uygulanıyor...")
    kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
    kmeans_labels = kmeans.fit_predict(X_scaled)
    print(f"✓ K-Means tamamlandı. Küme merkezleri: {kmeans.cluster_centers_.shape}\n")
    
    # Yöntem 2: Agglomerative Clustering
    print("Agglomerative Clustering uygulanıyor...")
    agglomerative = AgglomerativeClustering(n_clusters=n_clusters, linkage='ward')
    agglomerative_labels = agglomerative.fit_predict(X_scaled)
    print(f"✓ Agglomerative Clustering tamamlandı.\n")
    
    print("--- Kümeleme Tamamlandı ---\n")
    
    # DOCX kümeleme uygulaması
    add_heading_docx(doc, "4. KÜMELEME YÖNTEMLERİNİN UYGULANMASI", level=1)
    add_paragraph_docx(doc, "K-Means kümeleme algoritması uygulanmıştır. Parametreler: n_clusters=4, random_state=42, n_init=10")
    add_paragraph_docx(doc, "Agglomerative Clustering algoritması uygulanmıştır. Parametreler: n_clusters=4, linkage='ward'")

    # === 5. BAŞARI DEĞERLERİNİN HESAPLANMASI ===
    print_separator("5. BAŞARI DEĞERLERİNİN HESAPLANMASI")
    
    print("\n--- Başarı Metrikleri Hesaplanıyor ---\n")
    
    # K-Means metrikleri
    ari_kmeans = adjusted_rand_score(y_true, kmeans_labels)
    nmi_kmeans = normalized_mutual_info_score(y_true, kmeans_labels)
    completeness_kmeans = completeness_score(y_true, kmeans_labels)
    homogeneity_kmeans = homogeneity_score(y_true, kmeans_labels)
    silhouette_kmeans = silhouette_score(X_scaled, kmeans_labels)
    
    # Agglomerative metrikleri
    ari_agglomerative = adjusted_rand_score(y_true, agglomerative_labels)
    nmi_agglomerative = normalized_mutual_info_score(y_true, agglomerative_labels)
    completeness_agglomerative = completeness_score(y_true, agglomerative_labels)
    homogeneity_agglomerative = homogeneity_score(y_true, agglomerative_labels)
    silhouette_agglomerative = silhouette_score(X_scaled, agglomerative_labels)
    
    # Metrik açıklamaları
    print("Kullanılan Metrikler:")
    print("  - Adjusted Rand Index (ARI): Gerçek etiketlerle küme etiketlerinin uyumunu ölçer. [-1, 1] aralığında, 1 mükemmel uyum.")
    print("  - Normalized Mutual Information (NMI): İki kümeleme arasındaki normalleştirilmiş bilgi paylaşımını ölçer. [0, 1] aralığında.")
    print("  - Completeness Score: Her gerçek sınıfın tek bir kümede toplanma derecesi. [0, 1] aralığında.")
    print("  - Homogeneity Score: Her kümenin tek bir gerçek sınıftan oluşma derecesi. [0, 1] aralığında.")
    print("  - Silhouette Score: Gözetimsiz metrik. Küme içi sıkılık ve küme arası ayrışmayı ölçer. [-1, 1] aralığında.\n")
    
    # K-Means sonuçları
    print("--- K-Means Başarı Metrikleri ---")
    print(f"Adjusted Rand Index (ARI)        : {ari_kmeans:.4f}")
    print(f"Normalized Mutual Information    : {nmi_kmeans:.4f}")
    print(f"Completeness Score              : {completeness_kmeans:.4f}")
    print(f"Homogeneity Score               : {homogeneity_kmeans:.4f}")
    print(f"Silhouette Score                : {silhouette_kmeans:.4f}")
    
    # Agglomerative sonuçları
    print("\n--- Agglomerative Clustering Başarı Metrikleri ---")
    print(f"Adjusted Rand Index (ARI)        : {ari_agglomerative:.4f}")
    print(f"Normalized Mutual Information    : {nmi_agglomerative:.4f}")
    print(f"Completeness Score              : {completeness_agglomerative:.4f}")
    print(f"Homogeneity Score               : {homogeneity_agglomerative:.4f}")
    print(f"Silhouette Score                : {silhouette_agglomerative:.4f}\n")
    
    # DOCX başarı metrikleri
    add_heading_docx(doc, "5. BAŞARI DEĞERLERİNİN HESAPLANMASI", level=1)
    
    metric_explanation = """Adjusted Rand Index (ARI): Gerçek etiketlerle küme etiketlerinin uyumunu ölçer. [-1, 1] aralığında, 1 mükemmel uyum.
Normalized Mutual Information (NMI): İki kümeleme arasındaki normalleştirilmiş bilgi paylaşımını ölçer. [0, 1] aralığında.
Completeness Score: Her gerçek sınıfın tek bir kümede toplanma derecesi. [0, 1] aralığında.
Homogeneity Score: Her kümenin tek bir gerçek sınıftan oluşma derecesi. [0, 1] aralığında.
Silhouette Score: Gözetimsiz metrik. Küme içi sıkılık ve küme arası ayrışmayı ölçer. [-1, 1] aralığında."""
    
    add_paragraph_docx(doc, metric_explanation)
    
    # K-Means metrikleri tablosu
    add_heading_docx(doc, "K-Means Başarı Metrikleri", level=2)
    kmeans_metric_rows = [
        ['Adjusted Rand Index (ARI)', f"{ari_kmeans:.4f}"],
        ['Normalized Mutual Information (NMI)', f"{nmi_kmeans:.4f}"],
        ['Completeness Score', f"{completeness_kmeans:.4f}"],
        ['Homogeneity Score', f"{homogeneity_kmeans:.4f}"],
        ['Silhouette Score', f"{silhouette_kmeans:.4f}"]
    ]
    add_table_docx(doc, ['Metrik', 'Değer'], kmeans_metric_rows)
    
    # Agglomerative metrikleri tablosu
    add_heading_docx(doc, "Agglomerative Clustering Başarı Metrikleri", level=2)
    agg_metric_rows = [
        ['Adjusted Rand Index (ARI)', f"{ari_agglomerative:.4f}"],
        ['Normalized Mutual Information (NMI)', f"{nmi_agglomerative:.4f}"],
        ['Completeness Score', f"{completeness_agglomerative:.4f}"],
        ['Homogeneity Score', f"{homogeneity_agglomerative:.4f}"],
        ['Silhouette Score', f"{silhouette_agglomerative:.4f}"]
    ]
    add_table_docx(doc, ['Metrik', 'Değer'], agg_metric_rows)

    # === 6. YÖNTEMLERİN KARŞILAŞTIRILMASI ===
    print_separator("6. YÖNTEMLERİN KARŞILAŞTIRILMASI")
    
    print("\n--- Yöntemler Arası Karşılaştırma Tablosu ---\n")
    
    comparison_rows = [
        ['Adjusted Rand Index (ARI)', f"{ari_kmeans:.4f}", f"{ari_agglomerative:.4f}"],
        ['Normalized Mutual Information (NMI)', f"{nmi_kmeans:.4f}", f"{nmi_agglomerative:.4f}"],
        ['Completeness Score', f"{completeness_kmeans:.4f}", f"{completeness_agglomerative:.4f}"],
        ['Homogeneity Score', f"{homogeneity_kmeans:.4f}", f"{homogeneity_agglomerative:.4f}"],
        ['Silhouette Score', f"{silhouette_kmeans:.4f}", f"{silhouette_agglomerative:.4f}"]
    ]
    
    print(f"{'Metrik':<40} {'K-Means':<20} {'Agglomerative':<20}")
    print("-"*80)
    for row in comparison_rows:
        print(f"{row[0]:<40} {row[1]:<20} {row[2]:<20}")
    print("-"*80)
    
    # En iyi yöntemi belirle
    best_methods = {
        'ARI': 'K-Means' if ari_kmeans >= ari_agglomerative else 'Agglomerative',
        'NMI': 'K-Means' if nmi_kmeans >= nmi_agglomerative else 'Agglomerative',
        'Completeness': 'K-Means' if completeness_kmeans >= completeness_agglomerative else 'Agglomerative',
        'Homogeneity': 'K-Means' if homogeneity_kmeans >= homogeneity_agglomerative else 'Agglomerative',
        'Silhouette': 'K-Means' if silhouette_kmeans >= silhouette_agglomerative else 'Agglomerative'
    }
    
    print("\n--- En İyi Performans ---")
    for metric, method in best_methods.items():
        print(f"{metric}: {method}")
    
    # DOCX karşılaştırma tablosu
    add_heading_docx(doc, "6. YÖNTEMLERİN KARŞILAŞTIRILMASI", level=1)
    add_table_docx(doc, ['Metrik', 'K-Means', 'Agglomerative'], comparison_rows)
    
    # En iyi yöntemler
    add_heading_docx(doc, "En İyi Performans", level=2)
    best_methods_text = "\n".join([f"{metric}: {method}" for metric, method in best_methods.items()])
    add_paragraph_docx(doc, best_methods_text)

    # === 7. SONUÇ VE YORUMLAR ===
    print_separator("7. SONUÇ VE YORUMLAR")
    
    # Kümelerin dağılımı
    print("\n--- Küme Dağılımları ---\n")
    
    print("K-Means Küme Dağılımı:")
    kmeans_dist = pd.Series(kmeans_labels).value_counts().sort_index()
    for cluster_id, count in kmeans_dist.items():
        print(f"  Küme {cluster_id}: {count} örnek")
    
    print("\nAgglomerative Clustering Küme Dağılımı:")
    agg_dist = pd.Series(agglomerative_labels).value_counts().sort_index()
    for cluster_id, count in agg_dist.items():
        print(f"  Küme {cluster_id}: {count} örnek")
    
    # Gerçek sınıf dağılımı
    print("\nGerçek Sınıf Dağılımı:")
    true_dist = pd.Series(y_labels).value_counts()
    for class_name, count in true_dist.items():
        print(f"  {class_name}: {count} örnek")
    
    # Yorumlar
    print("\n--- Yorumlar ---\n")
    
    # Hangi metriklerde hangi yöntem daha iyi
    kmeans_wins = sum(1 for method in best_methods.values() if method == 'K-Means')
    agg_wins = sum(1 for method in best_methods.values() if method == 'Agglomerative')
    
    conclusion_text = f"""
    Her iki kümeleme yöntemi de veri seti üzerinde test edilmiştir. Toplam 5 metrik kullanılarak performans değerlendirmesi yapılmıştır:
    
    - K-Means, {kmeans_wins} metrikte en iyi performansı göstermiştir.
    - Agglomerative Clustering, {agg_wins} metrikte en iyi performansı göstermiştir.
    
    Adjusted Rand Index (ARI), gerçek etiketlerle küme etiketlerinin uyumunu ölçen en önemli metriktir. Bu metrikte 
    {'K-Means' if ari_kmeans >= ari_agglomerative else 'Agglomerative Clustering'} yöntemi daha yüksek skor elde etmiştir.
    
    Silhouette Score, gözetimsiz bir metrik olup gerçek etiketlere ihtiyaç duymaz. Bu metrik küme içi sıkılığı ve 
    küme arası ayrışmayı değerlendirir. Bu metrikte {'K-Means' if silhouette_kmeans >= silhouette_agglomerative else 'Agglomerative Clustering'} 
    yöntemi daha iyi performans göstermiştir.
    
    Genel olarak, her iki yöntem de benzer performans sergilemiştir. K-Means genellikle daha hızlı çalışırken, 
    Agglomerative Clustering hiyerarşik yapıyı koruma avantajına sahiptir.
    """
    
    print(conclusion_text)
    
    # DOCX sonuç ve yorumlar
    add_heading_docx(doc, "7. SONUÇ VE YORUMLAR", level=1)
    
    # Küme dağılımları tablosu
    add_heading_docx(doc, "Küme Dağılımları", level=2)
    cluster_dist_rows = []
    max_clusters = max(len(kmeans_dist), len(agg_dist))
    for i in range(max_clusters):
        kmeans_count = kmeans_dist.get(i, 0) if i in kmeans_dist.index else 0
        agg_count = agg_dist.get(i, 0) if i in agg_dist.index else 0
        cluster_dist_rows.append([f"Küme {i}", str(kmeans_count), str(agg_count)])
    add_table_docx(doc, ['Küme', 'K-Means', 'Agglomerative'], cluster_dist_rows)
    
    # Yorumlar
    add_heading_docx(doc, "Yorumlar", level=2)
    add_paragraph_docx(doc, conclusion_text.strip())
    
    # === KOD EKLEMESİ ===
    print_separator("8. KULLANILAN KOD")
    
    code_snippet = """
# Kümeleme yöntemlerinin uygulanması
from sklearn.cluster import KMeans, AgglomerativeClustering
from sklearn.metrics import adjusted_rand_score, silhouette_score, normalized_mutual_info_score

# K-Means
kmeans = KMeans(n_clusters=4, random_state=42, n_init=10)
kmeans_labels = kmeans.fit_predict(X_scaled)

# Agglomerative Clustering
agglomerative = AgglomerativeClustering(n_clusters=4, linkage='ward')
agglomerative_labels = agglomerative.fit_predict(X_scaled)

# Başarı metrikleri
ari_kmeans = adjusted_rand_score(y_true, kmeans_labels)
ari_agglomerative = adjusted_rand_score(y_true, agglomerative_labels)
silhouette_kmeans = silhouette_score(X_scaled, kmeans_labels)
silhouette_agglomerative = silhouette_score(X_scaled, agglomerative_labels)
"""
    
    print(code_snippet)
    
    # DOCX kod eklemesi
    add_heading_docx(doc, "8. KULLANILAN KOD", level=1)
    add_paragraph_docx(doc, code_snippet.strip(), italic=True)
    
    print_separator("RAPOR SONU")
    
    # DOCX dosyasını kaydet
    doc_filename = 'Rapor_Kumeleme_Analizi.docx'
    doc.save(doc_filename)
    print(f"\n{'='*80}")
    print(f"DOCX raporu oluşturuldu: {doc_filename}")
    print(f"{'='*80}\n")

except FileNotFoundError:
    error_msg = "HATA: 'Acoustic Features.csv' dosyası bulunamadı."
    print(error_msg)
    print("Lütfen dataset dosyasını kodla aynı dizine koyun.")
    add_paragraph_docx(doc, error_msg)
    doc.save('Rapor_Kumeleme_Analizi.docx')
except Exception as e:
    error_msg = f"Bir hata oluştu: {e}"
    print(error_msg)
    import traceback
    traceback.print_exc()
    add_paragraph_docx(doc, error_msg)
    doc.save('Rapor_Kumeleme_Analizi.docx')
