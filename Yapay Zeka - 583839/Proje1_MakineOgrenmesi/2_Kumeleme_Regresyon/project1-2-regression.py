# -*- coding: utf-8 -*-
import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LinearRegression
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import mean_squared_error, r2_score, mean_absolute_error
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
import sys
import io

# Windows terminalinde UTF-8 encoding sorununu coz
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# Uyarilari bastir
warnings.filterwarnings('ignore')

# DOCX rapor icin Document olustur
doc = Document()

# === DOCX YARDIMCI FONKSIYONLAR ===

def add_heading_docx(doc, text, level=1):
    """DOCX'e baslik ekle"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph_docx(doc, text, bold=False, italic=False):
    """DOCX'e paragraf ekle"""
    para = doc.add_paragraph(text)
    if bold:
        for run in para.runs:
            run.bold = True
    if italic:
        for run in para.runs:
            run.italic = True
    return para

def add_table_docx(doc, headers, rows):
    """DOCX'e tablo ekle"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Basliklari ekle
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Satirlari ekle
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    return table

def print_separator(title=""):
    """Bolum ayirici yazdir"""
    print("\n" + "="*80)
    if title:
        print(f" {title}")
        print("="*80)
    else:
        print("="*80)

# === RAPOR BASLANGIC ===
print_separator("REGRESYON (REGRESSION) ANALIZ RAPORU")
print("\nBoston Housing Dataset - Regresyon Yontemleri Karsilastirmasi\n")

# DOCX baslik
add_heading_docx(doc, "REGRESYON (REGRESSION) ANALIZ RAPORU", level=0)
add_paragraph_docx(doc, "Boston Housing Dataset - Regresyon Yontemleri Karsilastirmasi")
add_paragraph_docx(doc, "")

# === 1. REGRESYON NEDIR? ===
print_separator("1. REGRESYON TANIMI")

regression_definition = """Regresyon, makine ogrenmesinde gozetimli ogrenme (supervised learning) yontemlerinden biridir. Regresyon analizi, bir veya daha fazla bagimsiz degisken (ozellik) ile bagimli degisken (hedef) arasindaki iliskiyi modellemeye ve tahmin etmeye yarar. Siniflandirmadan farkli olarak, regresyon surekli sayisal degerler tahmin eder (ornegin: ev fiyati, sicaklik, maas). Regresyon modelleri, veri setindeki orneklerden ogrenip, yeni veriler icin tahminlerde bulunabilir."""

print("\n" + regression_definition + "\n")

# DOCX regresyon tanimi
add_heading_docx(doc, "1. REGRESYON TANIMI", level=1)
add_paragraph_docx(doc, regression_definition)

# === 2. REGRESYON YONTEMLERININ ACIKLAMASI ===
print_separator("2. KULLANILAN REGRESYON YONTEMLERI")

method_text_linear = """Dogrusal Regresyon (Linear Regression), en temel ve yaygin kullanilan regresyon yontemidir. Bagimsiz degiskenler ile bagimli degisken arasinda dogrusal bir iliski oldugunu varsayar. Model, en kucuk kareler yontemi (Ordinary Least Squares - OLS) kullanilarak egitilir ve hata karelerinin toplamini minimize eder. Formulu: y = b0 + b1*x1 + b2*x2 + ... + bn*xn seklindedir. Avantajlari: Basit, hizli, yorumlanabilir. Dezavantajlari: Dogrusal olmayan iliskilerde yetersiz kalir, aykiri degerlere duyarli."""

method_text_rf = """Rastgele Orman Regresyonu (Random Forest Regression), karar agaclarina dayali bir topluluk (ensemble) yontemidir. Birden fazla karar agaci olusturur ve her agacin tahminlerinin ortalamasini alarak nihai tahmini uretir. Her agac, veri setinin rastgele bir alt kumesi (bootstrap) ve ozelliklerin rastgele bir alt kumesi ile egitilir. Bu yontem, asiri uyum (overfitting) problemini azaltir ve dogrusal olmayan iliskileri yakalayabilir. Avantajlari: Yuksek dogruluk, ozellik onemi analizi, dogrusal olmayan iliskiler. Dezavantajlari: Hesaplama maliyeti yuksek, yorumlanmasi zor."""

print("\n2.1. Dogrusal Regresyon (Linear Regression)")
print("-"*80)
print(method_text_linear)
print("\n2.2. Rastgele Orman Regresyonu (Random Forest Regression)")
print("-"*80)
print(method_text_rf + "\n")

# DOCX yontem aciklamalari
add_heading_docx(doc, "2. KULLANILAN REGRESYON YONTEMLERI", level=1)
add_heading_docx(doc, "2.1. Dogrusal Regresyon (Linear Regression)", level=2)
add_paragraph_docx(doc, method_text_linear)
add_heading_docx(doc, "2.2. Rastgele Orman Regresyonu (Random Forest Regression)", level=2)
add_paragraph_docx(doc, method_text_rf)

# === 3. VERI SETI BILGISI ===
print_separator("3. VERI SETI BILGISI")

dataset_info = """
Veri Seti Adi: Boston Housing Dataset (California Housing Prices)
Kaynak: Scikit-learn kutuphanesi icin hazirlanmis standart veri seti
Alternatif Kaynak: https://www.kaggle.com/datasets/camnugent/california-housing-prices

Veri Seti Aciklamasi:
Bu veri seti, California eyaletindeki evlerin fiyatlarini tahmin etmek icin kullanilir. Veri seti, 1990 nufus sayimina dayanmaktadir. Her satir bir mahalle (block group) temsil eder ve ev fiyatlari ile ilgili cesitli ozellikler icerir.

Ozellikler:
- MedInc: Mahalle ortanca geliri (on binlerce dolar)
- HouseAge: Mahalle yas ortalamasi
- AveRooms: Ortalama oda sayisi
- AveBedrms: Ortalama yatak odasi sayisi
- Population: Mahalle nufusu
- AveOccup: Ortalama kisi sayisi
- Latitude: Enlem
- Longitude: Boylam

Hedef Degisken:
- MedHouseVal: Mahalle ortanca ev degeri (yuz binlerce dolar)
"""

print(dataset_info)

# DOCX veri seti bilgisi
add_heading_docx(doc, "3. VERI SETI BILGISI", level=1)
add_paragraph_docx(doc, dataset_info.strip())

# === 4. VERI YUKLEME VE HAZIRLAMA ===
print_separator("4. VERI YUKLEME VE HAZIRLAMA")

try:
    # California Housing veri setini yukle (scikit-learn'den)
    from sklearn.datasets import fetch_california_housing
    
    print("\nCalifornia Housing veri seti yukleniyor...")
    housing = fetch_california_housing()
    X = housing.data
    y = housing.target
    feature_names = housing.feature_names
    
    print(f"\nVeri Seti Yuklendi.")
    print(f"Ornek Sayisi: {X.shape[0]}, Ozellik Sayisi: {X.shape[1]}")
    print(f"Hedef Degisken: MedHouseVal (Ortanca Ev Degeri)")
    print(f"Ozellikler: {', '.join(feature_names)}\n")
    
    # Veriyi egitim ve test setlerine ayir
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
    
    print(f"Egitim Seti: {X_train.shape[0]} ornek")
    print(f"Test Seti: {X_test.shape[0]} ornek\n")
    
    # Veriyi normalize et (regresyon icin kritik)
    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train)
    X_test_scaled = scaler.transform(X_test)
    
    print("Veri normalize edildi (StandardScaler).\n")
    
    # DOCX veri bilgisi
    add_heading_docx(doc, "4. VERI YUKLEME VE HAZIRLAMA", level=1)
    add_paragraph_docx(doc, f"Ornek Sayisi: {X.shape[0]}, Ozellik Sayisi: {X.shape[1]}")
    add_paragraph_docx(doc, f"Hedef Degisken: MedHouseVal (Ortanca Ev Degeri)")
    add_paragraph_docx(doc, f"Ozellikler: {', '.join(feature_names)}")
    add_paragraph_docx(doc, f"Egitim Seti: {X_train.shape[0]} ornek, Test Seti: {X_test.shape[0]} ornek")
    add_paragraph_docx(doc, "Veri StandardScaler ile normalize edilmistir.")

    # === 5. REGRESYON YONTEMLERININ UYGULANMASI ===
    print_separator("5. REGRESYON YONTEMLERININ UYGULANMASI")
    
    print("--- Regresyon Yontemleri Uygulanıyor ---\n")
    
    # Yontem 1: Linear Regression
    print("Dogrusal Regresyon (Linear Regression) uygulanıyor...")
    lr_model = LinearRegression()
    lr_model.fit(X_train_scaled, y_train)
    y_pred_lr = lr_model.predict(X_test_scaled)
    print(f"✓ Dogrusal Regresyon tamamlandi.\n")
    
    # Yontem 2: Random Forest Regression
    print("Rastgele Orman Regresyonu (Random Forest Regression) uygulanıyor...")
    rf_model = RandomForestRegressor(n_estimators=100, random_state=42, max_depth=10)
    rf_model.fit(X_train_scaled, y_train)
    y_pred_rf = rf_model.predict(X_test_scaled)
    print(f"✓ Rastgele Orman Regresyonu tamamlandi.\n")
    
    print("--- Regresyon Tamamlandi ---\n")
    
    # DOCX regresyon uygulamasi
    add_heading_docx(doc, "5. REGRESYON YONTEMLERININ UYGULANMASI", level=1)
    add_paragraph_docx(doc, "Dogrusal Regresyon (Linear Regression) algoritmasi uygulanmistir.")
    add_paragraph_docx(doc, "Rastgele Orman Regresyonu (Random Forest Regression) algoritmasi uygulanmistir. Parametreler: n_estimators=100, random_state=42, max_depth=10")

    # === 6. BASARI METRIKLERININ ACIKLAMASI ===
    print_separator("6. BASARI METRIKLERININ ACIKLAMASI")
    
    metric_explanation_1 = """
6.1. Ortalama Karesel Hata (Mean Squared Error - MSE)

Formul: MSE = (1/n) * Σ(y_gercek - y_tahmin)²

Aciklama: MSE, gercek degerler ile tahmin edilen degerler arasindaki farkin karesinin ortalamasini alir. Hatalari karesi alinarak buyuk hatalara daha fazla agirlik verilir. Dusuk MSE degeri, modelin daha basarili oldugunu gosterir. MSE her zaman pozitif bir degerdir ve en iyi deger 0'dir (mukemmel tahmin).
"""

    metric_explanation_2 = """
6.2. Belirleme Katsayisi (R² Score - Coefficient of Determination)

Formul: R² = 1 - (SS_res / SS_tot)
        SS_res = Σ(y_gercek - y_tahmin)² (Hata Kareler Toplami)
        SS_tot = Σ(y_gercek - y_ortalama)² (Toplam Kareler Toplami)

Aciklama: R² skoru, modelin bagimli degiskendeki varyansın ne kadarini aciklayabildigini gosterir. Deger 0 ile 1 arasinda olup, 1'e yakin olmasi modelin cok basarili oldugunu gosterir. R² = 1, mukemmel tahmin anlamina gelir. R² = 0, modelin ortalama tahminden daha iyi olmadigi anlamina gelir. Negatif R² degerleri de mumkundur ve bu modelin cok kotu oldugunu gosterir.
"""
    
    print(metric_explanation_1)
    print(metric_explanation_2)
    
    # DOCX metrik aciklamalari
    add_heading_docx(doc, "6. BASARI METRIKLERININ ACIKLAMASI", level=1)
    
    add_heading_docx(doc, "6.1. Ortalama Karesel Hata (Mean Squared Error - MSE)", level=2)
    add_paragraph_docx(doc, "Formul: MSE = (1/n) * Σ(y_gercek - y_tahmin)²")
    add_paragraph_docx(doc, "Aciklama: MSE, gercek degerler ile tahmin edilen degerler arasindaki farkin karesinin ortalamasini alir. Hatalari karesi alinarak buyuk hatalara daha fazla agirlik verilir. Dusuk MSE degeri, modelin daha basarili oldugunu gosterir. MSE her zaman pozitif bir degerdir ve en iyi deger 0'dir (mukemmel tahmin).")
    
    add_heading_docx(doc, "6.2. Belirleme Katsayisi (R² Score - Coefficient of Determination)", level=2)
    add_paragraph_docx(doc, "Formul: R² = 1 - (SS_res / SS_tot)")
    add_paragraph_docx(doc, "        SS_res = Σ(y_gercek - y_tahmin)² (Hata Kareler Toplami)")
    add_paragraph_docx(doc, "        SS_tot = Σ(y_gercek - y_ortalama)² (Toplam Kareler Toplami)")
    add_paragraph_docx(doc, "Aciklama: R² skoru, modelin bagimli degiskendeki varyansın ne kadarini aciklayabildigini gosterir. Deger 0 ile 1 arasinda olup, 1'e yakin olmasi modelin cok basarili oldugunu gosterir. R² = 1, mukemmel tahmin anlamina gelir. R² = 0, modelin ortalama tahminden daha iyi olmadigi anlamina gelir.")

    # === 7. BASARI DEGERLERININ HESAPLANMASI ===
    print_separator("7. BASARI DEGERLERININ HESAPLANMASI")
    
    print("\n--- Basari Metrikleri Hesaplaniyor ---\n")
    
    # Linear Regression metrikleri
    mse_lr = mean_squared_error(y_test, y_pred_lr)
    r2_lr = r2_score(y_test, y_pred_lr)
    mae_lr = mean_absolute_error(y_test, y_pred_lr)
    rmse_lr = np.sqrt(mse_lr)
    
    # Random Forest metrikleri
    mse_rf = mean_squared_error(y_test, y_pred_rf)
    r2_rf = r2_score(y_test, y_pred_rf)
    mae_rf = mean_absolute_error(y_test, y_pred_rf)
    rmse_rf = np.sqrt(mse_rf)
    
    # Linear Regression sonuclari
    print("--- Dogrusal Regresyon (Linear Regression) Basari Metrikleri ---")
    print(f"Mean Squared Error (MSE)         : {mse_lr:.4f}")
    print(f"Root Mean Squared Error (RMSE)   : {rmse_lr:.4f}")
    print(f"Mean Absolute Error (MAE)        : {mae_lr:.4f}")
    print(f"R² Score (Coefficient of Det.)   : {r2_lr:.4f}")
    
    # Random Forest sonuclari
    print("\n--- Rastgele Orman Regresyonu (Random Forest) Basari Metrikleri ---")
    print(f"Mean Squared Error (MSE)         : {mse_rf:.4f}")
    print(f"Root Mean Squared Error (RMSE)   : {rmse_rf:.4f}")
    print(f"Mean Absolute Error (MAE)        : {mae_rf:.4f}")
    print(f"R² Score (Coefficient of Det.)   : {r2_rf:.4f}\n")
    
    # DOCX basari metrikleri
    add_heading_docx(doc, "7. BASARI DEGERLERININ HESAPLANMASI", level=1)
    
    # Linear Regression metrikleri tablosu
    add_heading_docx(doc, "Dogrusal Regresyon (Linear Regression) Basari Metrikleri", level=2)
    lr_metric_rows = [
        ['Mean Squared Error (MSE)', f"{mse_lr:.4f}"],
        ['Root Mean Squared Error (RMSE)', f"{rmse_lr:.4f}"],
        ['Mean Absolute Error (MAE)', f"{mae_lr:.4f}"],
        ['R² Score', f"{r2_lr:.4f}"]
    ]
    add_table_docx(doc, ['Metrik', 'Deger'], lr_metric_rows)
    
    # Random Forest metrikleri tablosu
    add_heading_docx(doc, "Rastgele Orman Regresyonu (Random Forest) Basari Metrikleri", level=2)
    rf_metric_rows = [
        ['Mean Squared Error (MSE)', f"{mse_rf:.4f}"],
        ['Root Mean Squared Error (RMSE)', f"{rmse_rf:.4f}"],
        ['Mean Absolute Error (MAE)', f"{mae_rf:.4f}"],
        ['R² Score', f"{r2_rf:.4f}"]
    ]
    add_table_docx(doc, ['Metrik', 'Deger'], rf_metric_rows)

    # === 8. YONTEMLERIN KARSILASTIRILMASI ===
    print_separator("8. YONTEMLERIN KARSILASTIRILMASI")
    
    print("\n--- Yontemler Arasi Karsilastirma Tablosu ---\n")
    
    comparison_rows = [
        ['Mean Squared Error (MSE)', f"{mse_lr:.4f}", f"{mse_rf:.4f}"],
        ['Root Mean Squared Error (RMSE)', f"{rmse_lr:.4f}", f"{rmse_rf:.4f}"],
        ['Mean Absolute Error (MAE)', f"{mae_lr:.4f}", f"{mae_rf:.4f}"],
        ['R² Score', f"{r2_lr:.4f}", f"{r2_rf:.4f}"]
    ]
    
    print(f"{'Metrik':<40} {'Linear Regression':<20} {'Random Forest':<20}")
    print("-"*80)
    for row in comparison_rows:
        print(f"{row[0]:<40} {row[1]:<20} {row[2]:<20}")
    print("-"*80)
    
    # En iyi yontemi belirle (MSE ve MAE icin dusuk iyi, R² icin yuksek iyi)
    best_methods = {
        'MSE': 'Linear Regression' if mse_lr <= mse_rf else 'Random Forest',
        'RMSE': 'Linear Regression' if rmse_lr <= rmse_rf else 'Random Forest',
        'MAE': 'Linear Regression' if mae_lr <= mae_rf else 'Random Forest',
        'R² Score': 'Linear Regression' if r2_lr >= r2_rf else 'Random Forest'
    }
    
    print("\n--- En Iyi Performans ---")
    for metric, method in best_methods.items():
        print(f"{metric}: {method}")
    
    # DOCX karsilastirma tablosu
    add_heading_docx(doc, "8. YONTEMLERIN KARSILASTIRILMASI", level=1)
    add_table_docx(doc, ['Metrik', 'Linear Regression', 'Random Forest'], comparison_rows)
    
    # En iyi yontemler
    add_heading_docx(doc, "En Iyi Performans", level=2)
    best_methods_text = "\n".join([f"{metric}: {method}" for metric, method in best_methods.items()])
    add_paragraph_docx(doc, best_methods_text)

    # === 9. SONUC VE YORUMLAR ===
    print_separator("9. SONUC VE YORUMLAR")
    
    print("\n--- Yorumlar ---\n")
    
    # Hangi metriklerde hangi yontem daha iyi
    lr_wins = sum(1 for method in best_methods.values() if method == 'Linear Regression')
    rf_wins = sum(1 for method in best_methods.values() if method == 'Random Forest')
    
    # Tahmin ornekleri
    print("--- Ornek Tahminler (Ilk 5 Test Orneği) ---\n")
    print(f"{'Gercek Deger':<15} {'LR Tahmini':<15} {'RF Tahmini':<15} {'LR Hatasi':<15} {'RF Hatasi':<15}")
    print("-"*75)
    for i in range(5):
        lr_error = abs(y_test[i] - y_pred_lr[i])
        rf_error = abs(y_test[i] - y_pred_rf[i])
        print(f"{y_test[i]:<15.4f} {y_pred_lr[i]:<15.4f} {y_pred_rf[i]:<15.4f} {lr_error:<15.4f} {rf_error:<15.4f}")
    print("-"*75)
    
    conclusion_text = f"""
Her iki regresyon yontemi de California Housing veri seti uzerinde test edilmistir. Toplam 4 metrik kullanilarak performans degerlendirmesi yapilmistir:

- Linear Regression, {lr_wins} metrikte en iyi performansi gostermistir.
- Random Forest Regression, {rf_wins} metrikte en iyi performansi gostermistir.

R² Score (Belirleme Katsayisi), regresyon modellerinin performansini degerlen diren en onemli metriktir. Bu metrikte {'Linear Regression' if r2_lr >= r2_rf else 'Random Forest Regression'} yontemi daha yuksek skor elde etmistir (LR: {r2_lr:.4f}, RF: {r2_rf:.4f}).

Mean Squared Error (MSE), tahminlerdeki ortalama hatayi gosterir. Dusuk MSE degeri daha iyi performans anlamina gelir. Bu metrikte {'Linear Regression' if mse_lr <= mse_rf else 'Random Forest Regression'} yontemi daha basarili olmustur (LR: {mse_lr:.4f}, RF: {mse_rf:.4f}).

Genel Degerlendirme:
{'Random Forest Regression yontemi, dogrusal olmayan iliskileri daha iyi yakalayarak genel olarak daha yuksek basari gostermistir. Ancak Linear Regression, basitligi ve yorumlanabilirligi acisindan avantajlidir.' if rf_wins >= lr_wins else 'Linear Regression yontemi, veri setindeki iliskileri basariyla modelleyerek iyi performans gostermistir. Bu veri setinde dogrusal iliski baskın olabilir.'}

Her iki model de kabul edilebilir R² skorlarina sahiptir ({r2_lr:.4f} ve {r2_rf:.4f}), bu da modellerin veri setindeki varyansın onemli bir kismini aciklayabildigini gosterir.
"""
    
    print(conclusion_text)
    
    # DOCX sonuc ve yorumlar
    add_heading_docx(doc, "9. SONUC VE YORUMLAR", level=1)
    
    # Ornek tahminler tablosu
    add_heading_docx(doc, "Ornek Tahminler (Ilk 5 Test Ornegi)", level=2)
    sample_rows = []
    for i in range(5):
        lr_error = abs(y_test[i] - y_pred_lr[i])
        rf_error = abs(y_test[i] - y_pred_rf[i])
        sample_rows.append([
            f"{y_test[i]:.4f}",
            f"{y_pred_lr[i]:.4f}",
            f"{y_pred_rf[i]:.4f}",
            f"{lr_error:.4f}",
            f"{rf_error:.4f}"
        ])
    add_table_docx(doc, ['Gercek Deger', 'LR Tahmini', 'RF Tahmini', 'LR Hatasi', 'RF Hatasi'], sample_rows)
    
    # Yorumlar
    add_heading_docx(doc, "Yorumlar", level=2)
    add_paragraph_docx(doc, conclusion_text.strip())
    
    # === 10. KULLANILAN KOD ===
    print_separator("10. KULLANILAN KOD")
    
    code_snippet = """
# Regresyon yontemlerinin uygulanmasi
from sklearn.linear_model import LinearRegression
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import mean_squared_error, r2_score, mean_absolute_error

# Veriyi egitim ve test setlerine ayir
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Veriyi normalize et
scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)

# Linear Regression
lr_model = LinearRegression()
lr_model.fit(X_train_scaled, y_train)
y_pred_lr = lr_model.predict(X_test_scaled)

# Random Forest Regression
rf_model = RandomForestRegressor(n_estimators=100, random_state=42, max_depth=10)
rf_model.fit(X_train_scaled, y_train)
y_pred_rf = rf_model.predict(X_test_scaled)

# Basari metrikleri
mse_lr = mean_squared_error(y_test, y_pred_lr)
r2_lr = r2_score(y_test, y_pred_lr)
mse_rf = mean_squared_error(y_test, y_pred_rf)
r2_rf = r2_score(y_test, y_pred_rf)
"""
    
    print(code_snippet)
    
    # DOCX kod eklemesi
    add_heading_docx(doc, "10. KULLANILAN KOD", level=1)
    add_paragraph_docx(doc, code_snippet.strip(), italic=True)
    
    print_separator("RAPOR SONU")
    
    # DOCX dosyasini kaydet
    doc_filename = 'Rapor_Regresyon_Analizi.docx'
    doc.save(doc_filename)
    print(f"\n{'='*80}")
    print(f"DOCX raporu olusturuldu: {doc_filename}")
    print(f"{'='*80}\n")

except Exception as e:
    error_msg = f"Bir hata olustu: {e}"
    print(error_msg)
    import traceback
    traceback.print_exc()
    add_paragraph_docx(doc, error_msg)
    doc.save('Rapor_Regresyon_Analizi.docx')
