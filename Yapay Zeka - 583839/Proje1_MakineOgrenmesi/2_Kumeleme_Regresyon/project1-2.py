# -*- coding: utf-8 -*-
import pandas as pd
import numpy as np
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.cluster import KMeans, AgglomerativeClustering
from sklearn.metrics import adjusted_rand_score, silhouette_score, normalized_mutual_info_score, completeness_score, homogeneity_score
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
import sys
import io

# Windows terminalinde UTF-8 encoding sorununu coz
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# Uyarilari bastir
warnings.filterwarnings('ignore')

# DOCX rapor icin Document olustur
doc = Document()

# === DOCX YARDIMCI FONKSIYONLAR ===

def add_heading_docx(doc, text, level=1):
    """DOCX'e baslik ekle"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph_docx(doc, text, bold=False, italic=False):
    """DOCX'e paragraf ekle"""
    para = doc.add_paragraph(text)
    if bold:
        for run in para.runs:
            run.bold = True
    if italic:
        for run in para.runs:
            run.italic = True
    return para

def add_table_docx(doc, headers, rows):
    """DOCX'e tablo ekle"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Basliklari ekle
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Satirlari ekle
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    return table

def print_separator(title=""):
    """Bolum ayirici yazdir"""
    print("\n" + "="*80)
    if title:
        print(f" {title}")
        print("="*80)
    else:
        print("="*80)

# === RAPOR BASLANGIC ===
print_separator("KUMELEME (CLUSTERING) ANALIZ RAPORU")
print("\nAcoustic Features Dataset - Kumeleme Yontemleri Karsilastirmasi\n")

# DOCX baslik
add_heading_docx(doc, "KUMELEME (CLUSTERING) ANALIZ RAPORU", level=0)
add_paragraph_docx(doc, "Acoustic Features Dataset - Kumeleme Yontemleri Karsilastirmasi")
add_paragraph_docx(doc, "")

# === 1. PROBLEM ACIKLAMASI ===
print_separator("1. PROBLEM ACIKLAMASI")

problem_text = """Bu calismada, muzik parcalarinin akustik ozelliklerine gore kumeleme (clustering) analizi yapilmaktadir. Kumeleme, goz etimsiz ogrenme (unsupervised learning) yontemidir ve verideki dogal gruplari kesfetmeye calisir. Gercek etiketler bilindigi icin (angry, happy, relax, sad), kumeleme yontemlerinin performansini gercek etiketlerle karsilastirarak degerlendirebiliriz. Iki farkli kumeleme algoritmasi (K-Means ve Agglomerative Clustering) kullanilarak veri kumelere ayrilmis ve basari metrikleri ile karsilastirilmistir."""

print("\n" + problem_text + "\n")

# DOCX problem aciklamasi
add_heading_docx(doc, "1. PROBLEM ACIKLAMASI", level=1)
add_paragraph_docx(doc, problem_text)

# === 2. KUMELEME YONTEMLERININ ACIKLAMASI ===
print_separator("2. KULLANILAN KUMELEME YONTEMLERI")

method_text_kmeans = """K-Means, en populer kumeleme algoritmalarindan biridir. Algoritma, veriyi k adet kumeye ayirmaya calisir. Isleyisi su sekildedir: (1) k adet rastgele merkez (centroid) secilir, (2) Her veri noktasi en yakin merkeze atanir, (3) Merkezler, atanan noktalarin ortalamasi olarak guncellenir, (4) Adimlar 2-3 iterasyon olarak tekrarlanir ve kume merkezleri sabitlenince durur. Avantajlari: Hizli ve olceklenebilir, basit ve anlasilir. Dezavantajlari: Kume sayisinin onceden belirlenmesi gerekir, kuresel olmayan kumelerde zorlanir."""

method_text_agg = """Agglomerative Clustering (Hiyerarsik Kumeleme), hiyerarsik bir kumeleme yontemidir. Algoritma, her veri noktasini baslangicta ayri bir kume olarak ele alir ve en yakin kumeleri iteratif olarak birlestirir (bottom-up yaklasim). Ward linkage kriteri kullanildiginda, birlestirme islemi kumeler ici varyansini minimize edecek sekilde yapilir. Avantajlari: Kume sayisini sonradan belirleme esnekligi, dendrogram ile gorsellestirme imkani. Dezavantajlari: Hesaplama maliyeti yuksek (O(n³)), buyuk veri setlerinde yavas."""

print("\n2.1. K-Means Kumeleme")
print("-"*80)
print(method_text_kmeans)
print("\n2.2. Agglomerative Clustering (Hiyerarsik Kumeleme)")
print("-"*80)
print(method_text_agg + "\n")

# DOCX yontem aciklamalari
add_heading_docx(doc, "2. KULLANILAN KUMELEME YONTEMLERI", level=1)
add_heading_docx(doc, "2.1. K-Means Kumeleme", level=2)
add_paragraph_docx(doc, method_text_kmeans)
add_heading_docx(doc, "2.2. Agglomerative Clustering (Hiyerarsik Kumeleme)", level=2)
add_paragraph_docx(doc, method_text_agg)

# === 3. VERI YUKLEME VE HAZIRLAMA ===
print_separator("3. VERI YUKLEME VE HAZIRLAMA")

try:
    # Veri setini yukle (Acoustic Features.csv kullaniliyor)
    df = pd.read_csv('Acoustic Features.csv')
    
    # Class sutununu etiketler (y_true) olarak ayir
    y_labels = df['Class'].values
    
    # Class sutunu disindaki tum sutunlari ozellikler (X) olarak al
    X = df.drop('Class', axis=1).values
    
    # Etiketleri encode et
    le = LabelEncoder()
    y_true = le.fit_transform(y_labels)
    class_names = le.classes_
    
    # Bilinen sinif sayisi
    n_clusters = len(class_names)
    
    print(f"\nVeri Seti Yuklendi.")
    print(f"Ornek Sayisi: {X.shape[0]}, Ozellik Sayisi: {X.shape[1]}")
    print(f"Kume Sayisi: {n_clusters}")
    print(f"Siniflar: {class_names}\n")
    
    # Veriyi normalize et (kumeleme icin kritik)
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)
    
    print("Veri normalize edildi (StandardScaler).\n")
    
    # DOCX veri bilgisi
    add_heading_docx(doc, "3. VERI YUKLEME VE HAZIRLAMA", level=1)
    add_paragraph_docx(doc, f"Ornek Sayisi: {X.shape[0]}, Ozellik Sayisi: {X.shape[1]}")
    add_paragraph_docx(doc, f"Kume Sayisi: {n_clusters}")
    add_paragraph_docx(doc, f"Siniflar: {', '.join(class_names)}")
    add_paragraph_docx(doc, "Veri StandardScaler ile normalize edilmistir.")

    # === 4. KUMELEME YONTEMLERININ UYGULANMASI ===
    print_separator("4. KUMELEME YONTEMLERININ UYGULANMASI")
    
    print("--- Kumeleme Yontemleri Uygulanıyor ---\n")
    
    # Yontem 1: K-Means
    print("K-Means kumeleme uygulanıyor...")
    kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
    kmeans_labels = kmeans.fit_predict(X_scaled)
    print(f"✓ K-Means tamamlandi. Kume merkezleri: {kmeans.cluster_centers_.shape}\n")
    
    # Yontem 2: Agglomerative Clustering
    print("Agglomerative Clustering uygulanıyor...")
    agglomerative = AgglomerativeClustering(n_clusters=n_clusters, linkage='ward')
    agglomerative_labels = agglomerative.fit_predict(X_scaled)
    print(f"✓ Agglomerative Clustering tamamlandi.\n")
    
    print("--- Kumeleme Tamamlandi ---\n")
    
    # DOCX kumeleme uygulamasi
    add_heading_docx(doc, "4. KUMELEME YONTEMLERININ UYGULANMASI", level=1)
    add_paragraph_docx(doc, "K-Means kumeleme algoritmasi uygulanmistir. Parametreler: n_clusters=4, random_state=42, n_init=10")
    add_paragraph_docx(doc, "Agglomerative Clustering algoritmasi uygulanmistir. Parametreler: n_clusters=4, linkage='ward'")

    # === 5. BASARI DEGERLERININ HESAPLANMASI ===
    print_separator("5. BASARI DEGERLERININ HESAPLANMASI")
    
    print("\n--- Basari Metrikleri Hesaplaniyor ---\n")
    
    # K-Means metrikleri
    ari_kmeans = adjusted_rand_score(y_true, kmeans_labels)
    nmi_kmeans = normalized_mutual_info_score(y_true, kmeans_labels)
    completeness_kmeans = completeness_score(y_true, kmeans_labels)
    homogeneity_kmeans = homogeneity_score(y_true, kmeans_labels)
    silhouette_kmeans = silhouette_score(X_scaled, kmeans_labels)
    
    # Agglomerative metrikleri
    ari_agglomerative = adjusted_rand_score(y_true, agglomerative_labels)
    nmi_agglomerative = normalized_mutual_info_score(y_true, agglomerative_labels)
    completeness_agglomerative = completeness_score(y_true, agglomerative_labels)
    homogeneity_agglomerative = homogeneity_score(y_true, agglomerative_labels)
    silhouette_agglomerative = silhouette_score(X_scaled, agglomerative_labels)
    
    # Metrik aciklamalari
    print("Kullanilan Metrikler:")
    print("  - Adjusted Rand Index (ARI): Gercek etiketlerle kume etiketlerinin uyumunu olcer. [-1, 1] araliginda, 1 mukemmel uyum.")
    print("  - Normalized Mutual Information (NMI): Iki kumeleme arasindaki normallestirilmis bilgi paylasimini olcer. [0, 1] araliginda.")
    print("  - Completeness Score: Her gercek sinifin tek bir kumede toplanma derecesi. [0, 1] araliginda.")
    print("  - Homogeneity Score: Her kumenin tek bir gercek siniftan olusma derecesi. [0, 1] araliginda.")
    print("  - Silhouette Score: Gozet imsiz metrik. Kume ici sikilik ve kume arasi ayrismavi olcer. [-1, 1] araliginda.\n")
    
    # K-Means sonuclari
    print("--- K-Means Basari Metrikleri ---")
    print(f"Adjusted Rand Index (ARI)        : {ari_kmeans:.4f}")
    print(f"Normalized Mutual Information    : {nmi_kmeans:.4f}")
    print(f"Completeness Score              : {completeness_kmeans:.4f}")
    print(f"Homogeneity Score               : {homogeneity_kmeans:.4f}")
    print(f"Silhouette Score                : {silhouette_kmeans:.4f}")
    
    # Agglomerative sonuclari
    print("\n--- Agglomerative Clustering Basari Metrikleri ---")
    print(f"Adjusted Rand Index (ARI)        : {ari_agglomerative:.4f}")
    print(f"Normalized Mutual Information    : {nmi_agglomerative:.4f}")
    print(f"Completeness Score              : {completeness_agglomerative:.4f}")
    print(f"Homogeneity Score               : {homogeneity_agglomerative:.4f}")
    print(f"Silhouette Score                : {silhouette_agglomerative:.4f}\n")
    
    # DOCX basari metrikleri
    add_heading_docx(doc, "5. BASARI DEGERLERININ HESAPLANMASI", level=1)
    
    metric_explanation = """Adjusted Rand Index (ARI): Gercek etiketlerle kume etiketlerinin uyumunu olcer. [-1, 1] araliginda, 1 mukemmel uyum.
Normalized Mutual Information (NMI): Iki kumeleme arasindaki normallestirilmis bilgi paylasimini olcer. [0, 1] araliginda.
Completeness Score: Her gercek sinifin tek bir kumede toplanma derecesi. [0, 1] araliginda.
Homogeneity Score: Her kumenin tek bir gercek siniftan olusma derecesi. [0, 1] araliginda.
Silhouette Score: Gozet imsiz metrik. Kume ici sikilik ve kume arasi ayrismavi olcer. [-1, 1] araliginda."""
    
    add_paragraph_docx(doc, metric_explanation)
    
    # K-Means metrikleri tablosu
    add_heading_docx(doc, "K-Means Basari Metrikleri", level=2)
    kmeans_metric_rows = [
        ['Adjusted Rand Index (ARI)', f"{ari_kmeans:.4f}"],
        ['Normalized Mutual Information (NMI)', f"{nmi_kmeans:.4f}"],
        ['Completeness Score', f"{completeness_kmeans:.4f}"],
        ['Homogeneity Score', f"{homogeneity_kmeans:.4f}"],
        ['Silhouette Score', f"{silhouette_kmeans:.4f}"]
    ]
    add_table_docx(doc, ['Metrik', 'Deger'], kmeans_metric_rows)
    
    # Agglomerative metrikleri tablosu
    add_heading_docx(doc, "Agglomerative Clustering Basari Metrikleri", level=2)
    agg_metric_rows = [
        ['Adjusted Rand Index (ARI)', f"{ari_agglomerative:.4f}"],
        ['Normalized Mutual Information (NMI)', f"{nmi_agglomerative:.4f}"],
        ['Completeness Score', f"{completeness_agglomerative:.4f}"],
        ['Homogeneity Score', f"{homogeneity_agglomerative:.4f}"],
        ['Silhouette Score', f"{silhouette_agglomerative:.4f}"]
    ]
    add_table_docx(doc, ['Metrik', 'Deger'], agg_metric_rows)

    # === 6. YONTEMLERIN KARSILASTIRILMASI ===
    print_separator("6. YONTEMLERIN KARSILASTIRILMASI")
    
    print("\n--- Yontemler Arasi Karsilastirma Tablosu ---\n")
    
    comparison_rows = [
        ['Adjusted Rand Index (ARI)', f"{ari_kmeans:.4f}", f"{ari_agglomerative:.4f}"],
        ['Normalized Mutual Information (NMI)', f"{nmi_kmeans:.4f}", f"{nmi_agglomerative:.4f}"],
        ['Completeness Score', f"{completeness_kmeans:.4f}", f"{completeness_agglomerative:.4f}"],
        ['Homogeneity Score', f"{homogeneity_kmeans:.4f}", f"{homogeneity_agglomerative:.4f}"],
        ['Silhouette Score', f"{silhouette_kmeans:.4f}", f"{silhouette_agglomerative:.4f}"]
    ]
    
    print(f"{'Metrik':<40} {'K-Means':<20} {'Agglomerative':<20}")
    print("-"*80)
    for row in comparison_rows:
        print(f"{row[0]:<40} {row[1]:<20} {row[2]:<20}")
    print("-"*80)
    
    # En iyi yontemi belirle
    best_methods = {
        'ARI': 'K-Means' if ari_kmeans >= ari_agglomerative else 'Agglomerative',
        'NMI': 'K-Means' if nmi_kmeans >= nmi_agglomerative else 'Agglomerative',
        'Completeness': 'K-Means' if completeness_kmeans >= completeness_agglomerative else 'Agglomerative',
        'Homogeneity': 'K-Means' if homogeneity_kmeans >= homogeneity_agglomerative else 'Agglomerative',
        'Silhouette': 'K-Means' if silhouette_kmeans >= silhouette_agglomerative else 'Agglomerative'
    }
    
    print("\n--- En Iyi Performans ---")
    for metric, method in best_methods.items():
        print(f"{metric}: {method}")
    
    # DOCX karsilastirma tablosu
    add_heading_docx(doc, "6. YONTEMLERIN KARSILASTIRILMASI", level=1)
    add_table_docx(doc, ['Metrik', 'K-Means', 'Agglomerative'], comparison_rows)
    
    # En iyi yontemler
    add_heading_docx(doc, "En Iyi Performans", level=2)
    best_methods_text = "\n".join([f"{metric}: {method}" for metric, method in best_methods.items()])
    add_paragraph_docx(doc, best_methods_text)

    # === 7. SONUC VE YORUMLAR ===
    print_separator("7. SONUC VE YORUMLAR")
    
    # Kumelerin dagilimi
    print("\n--- Kume Dagilimlari ---\n")
    
    print("K-Means Kume Dagilimi:")
    kmeans_dist = pd.Series(kmeans_labels).value_counts().sort_index()
    for cluster_id, count in kmeans_dist.items():
        print(f"  Kume {cluster_id}: {count} ornek")
    
    print("\nAgglomerative Clustering Kume Dagilimi:")
    agg_dist = pd.Series(agglomerative_labels).value_counts().sort_index()
    for cluster_id, count in agg_dist.items():
        print(f"  Kume {cluster_id}: {count} ornek")
    
    # Gercek sinif dagilimi
    print("\nGercek Sinif Dagilimi:")
    true_dist = pd.Series(y_labels).value_counts()
    for class_name, count in true_dist.items():
        print(f"  {class_name}: {count} ornek")
    
    # Yorumlar
    print("\n--- Yorumlar ---\n")
    
    # Hangi metriklerde hangi yontem daha iyi
    kmeans_wins = sum(1 for method in best_methods.values() if method == 'K-Means')
    agg_wins = sum(1 for method in best_methods.values() if method == 'Agglomerative')
    
    conclusion_text = f"""
    Her iki kumeleme yontemi de veri seti uzerinde test edilmistir. Toplam 5 metrik kullanilarak performans degerlendirmesi yapilmistir:
    
    - K-Means, {kmeans_wins} metrikte en iyi performansi gostermistir.
    - Agglomerative Clustering, {agg_wins} metrikte en iyi performansi gostermistir.
    
    Adjusted Rand Index (ARI), gercek etiketlerle kume etiketlerinin uyumunu olcen en onemli metriktir. Bu metrikte 
    {'K-Means' if ari_kmeans >= ari_agglomerative else 'Agglomerative Clustering'} yontemi daha yuksek skor elde etmistir.
    
    Silhouette Score, gozetimsiz bir metrik olup gercek etiketlere ihtiyac duymaz. Bu metrik kume ici sikiligi ve 
    kume arasi ayrismayi degerlendirir. Bu metrikte {'K-Means' if silhouette_kmeans >= silhouette_agglomerative else 'Agglomerative Clustering'} 
    yontemi daha iyi performans gostermistir.
    
    Genel olarak, her iki yontem de benzer performans sergilemistir. K-Means genellikle daha hizli calisirken, 
    Agglomerative Clustering hiyerarsik yapiyi koruma avantajina sahiptir.
    """
    
    print(conclusion_text)
    
    # DOCX sonuc ve yorumlar
    add_heading_docx(doc, "7. SONUC VE YORUMLAR", level=1)
    
    # Kume dagilimlari tablosu
    add_heading_docx(doc, "Kume Dagilimlari", level=2)
    cluster_dist_rows = []
    max_clusters = max(len(kmeans_dist), len(agg_dist))
    for i in range(max_clusters):
        kmeans_count = kmeans_dist.get(i, 0) if i in kmeans_dist.index else 0
        agg_count = agg_dist.get(i, 0) if i in agg_dist.index else 0
        cluster_dist_rows.append([f"Kume {i}", str(kmeans_count), str(agg_count)])
    add_table_docx(doc, ['Kume', 'K-Means', 'Agglomerative'], cluster_dist_rows)
    
    # Yorumlar
    add_heading_docx(doc, "Yorumlar", level=2)
    add_paragraph_docx(doc, conclusion_text.strip())
    
    # === KOD EKLEMESI ===
    print_separator("8. KULLANILAN KOD")
    
    code_snippet = """
# Kumeleme yontemlerinin uygulanmasi
from sklearn.cluster import KMeans, AgglomerativeClustering
from sklearn.metrics import adjusted_rand_score, silhouette_score, normalized_mutual_info_score

# K-Means
kmeans = KMeans(n_clusters=4, random_state=42, n_init=10)
kmeans_labels = kmeans.fit_predict(X_scaled)

# Agglomerative Clustering
agglomerative = AgglomerativeClustering(n_clusters=4, linkage='ward')
agglomerative_labels = agglomerative.fit_predict(X_scaled)

# Basari metrikleri
ari_kmeans = adjusted_rand_score(y_true, kmeans_labels)
ari_agglomerative = adjusted_rand_score(y_true, agglomerative_labels)
silhouette_kmeans = silhouette_score(X_scaled, kmeans_labels)
silhouette_agglomerative = silhouette_score(X_scaled, agglomerative_labels)
"""
    
    print(code_snippet)
    
    # DOCX kod eklemesi
    add_heading_docx(doc, "8. KULLANILAN KOD", level=1)
    add_paragraph_docx(doc, code_snippet.strip(), italic=True)
    
    print_separator("RAPOR SONU")
    
    # DOCX dosyasini kaydet
    doc_filename = 'Rapor_Kumeleme_Analizi.docx'
    doc.save(doc_filename)
    print(f"\n{'='*80}")
    print(f"DOCX raporu olusturuldu: {doc_filename}")
    print(f"{'='*80}\n")

except FileNotFoundError:
    error_msg = "HATA: 'Acoustic Features.csv' dosyasi bulunamadi."
    print(error_msg)
    print("Lutfen dataset dosyasini kodla ayni dizine koyun.")
    add_paragraph_docx(doc, error_msg)
    doc.save('Rapor_Kumeleme_Analizi.docx')
except Exception as e:
    error_msg = f"Bir hata olustu: {e}"
    print(error_msg)
    import traceback
    traceback.print_exc()
    add_paragraph_docx(doc, error_msg)
    doc.save('Rapor_Kumeleme_Analizi.docx')
